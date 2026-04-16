from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("/Users/wnluo/Desktop/code/rust_waf/paper/基于 Rust 的 L4∕L7 协同 Web 防护系统设计与实现.docx")
BACKUP = DOCX.with_name("基于 Rust 的 L4∕L7 协同 Web 防护系统设计与实现_before_literature_support.docx")


INSERTIONS = [
    (
        "本文的研究方法主要包括四类",
        [
            "从文献依据看，网络应用防火墙并不是单纯的访问控制列表，而是位于客户端与Web应用之间、能够结合HTTP语义、规则策略和代理转发机制进行检测与处置的安全组件[1]。OWASP关于ModSecurity和Core Rule Set的资料表明，开源WAF通常由检测引擎、规则集合、动作处置和日志审计等部分共同构成[2][3]；WAFEC项目也强调，评价WAF能力时不能只关注是否能够拦截某一类攻击，还应考察部署模式、协议兼容性、规则维护、误报控制和运行观测等指标[15]。因此，本文将研究对象界定为“检测、治理、观测和管理”相结合的综合性Web防护系统，而不是仅对单一拦截函数或单个规则库进行说明。",
        ],
    ),
    (
        "在系统实现语言方面，Rust生态的成熟为安全系统提供了新的实现选择",
        [
            "综合已有资料可以看出，国外开源实践更强调规则体系、协议适配和安全评估框架，相关标准文档则持续推动HTTP语义、HTTP/2、HTTP/3、QUIC和TLS等基础协议演进[7][8][9][10][11][12]。这些成果为本文系统设计提供了两类启示：一方面，防护系统必须尽量贴近协议标准，避免只依赖字符串级过滤；另一方面，工程实现需要把规则检测、连接治理、事件审计和配置管理放在同一架构中考虑。本文后续章节对四层/七层协同、统一请求抽象和控制台管理的分析，正是在上述研究脉络下展开。",
        ],
    ),
    (
        "两者并非替代关系，而应形成“先粗筛、后精判”的协同关系。",
        [
            "这种分层思想与现有网络安全工程实践具有一致性。应用层规则能够表达更丰富的攻击语义，但其前提是系统已经完成协议解析、请求重组和字段抽取，计算成本相对较高；连接层或传输层判定虽然语义较弱，却能够在更早阶段完成速率限制、异常连接识别和资源保护。WAFEC关于部署与检测能力的讨论也提示，WAF系统的有效性取决于检测深度、处置位置和运行成本之间的平衡[15]。因此，本文采用四层快速治理与七层精细识别相结合的方式，以降低高频异常流量对后续解析链路的影响。",
        ],
    ),
    (
        "如果为每一种协议单独编写完整检测逻辑，不仅会造成代码重复，也会使规则表达与管理复杂化。",
        [
            "从协议标准角度看，HTTP/1.1更接近传统文本请求模型，而HTTP/2通过二进制分帧、多路复用和头部压缩改变了请求承载方式，HTTP/3又进一步建立在QUIC之上，使连接标识、握手过程和传输层行为发生变化[7][8][9][12]。如果防护系统只面向某一种协议编写规则，就容易在协议升级或网关转发场景下出现检测盲区。统一请求抽象的作用，是把协议差异限制在接入转换层，使规则引擎和行为检测模块面对稳定的数据结构，从而提高系统的可维护性和可扩展性。",
        ],
    ),
    (
        "若规则命中，系统根据动作类型执行放行、阻断、告警记录或返回自定义HTTP响应。",
        [
            "从相关安全实践看，规则引擎的核心价值在于将攻击特征、处置动作和审计记录进行结构化表达。ModSecurity及其规则集的实践说明，通用规则能够覆盖大量常见攻击模式，但在具体业务中仍需要根据站点路径、请求方法、参数结构和响应策略进行调整[2][3]。因此，本文系统没有把规则匹配简单理解为固定字符串比较，而是将规则启停、层级归属、动作类型、自定义响应和事件记录一并纳入设计，使规则配置能够服务于后续运维分析和策略复盘。",
        ],
    ),
    (
        "该公式表明，在给定存储空间下，Bloom Filter可以通过牺牲少量误判概率换取更低的集合查询开销。",
        [
            "需要强调的是，Bloom Filter的误判特性决定了它更适合作为预筛工具，而不能直接作为最终安全裁决依据[6]。在Web防护场景中，如果将“可能命中”直接解释为“必须阻断”，可能会放大误报并影响正常业务访问；更合理的做法是将其放在黑名单、特征集合或热点对象判断的前置阶段，再交由精确规则、行为评分或人工配置进行复核。本文系统采用“快速预判+精确校验”的描述，正是为了在效率和可靠性之间保持平衡。",
        ],
    ),
    (
        "每一次阻断、挑战或告警都会被写入事件与画像数据，再反向影响后续分桶风险、行为分值和前端展示。",
        [
            "从学术表述上看，本文所谓“核心算法”并不等同于某一个孤立数学公式，而是指贯穿连接接入、协议抽象、风险评分、规则裁决和事件反馈的决策流程。已有WAF资料和Web安全实践表明，实际防护效果往往取决于规则、上下文、阈值、日志和运维处置之间的协同关系[1][14][15]。因此，本节在介绍具体公式和伪代码时，重点说明算法输入、状态维护、阈值比较、动作输出和前端可解释展示之间的关系，使系统实现能够与前文的理论基础和需求分析形成闭环。",
        ],
    ),
]


def set_body_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(24)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(20)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Times New Roman")
        run.font.size = Pt(12)
        run.bold = False


def add_citation_runs(paragraph, text):
    cursor = 0
    for match in re.finditer(r"\[(?:\d+)(?:\]\[\d+)*\]", text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        parts = re.findall(r"\[\d+\]", token)
        for part in parts:
            run = paragraph.add_run(part)
            run.font.superscript = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])
    set_body_style(paragraph)


def insert_after(paragraph, new_text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p = new_p
    inserted._element = new_p
    add_citation_runs(inserted, new_text)
    return inserted


def main():
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    existing_text = "\n".join(p.text for p in doc.paragraphs)
    inserted_count = 0

    for anchor, paragraphs in INSERTIONS:
        if any(paragraphs[0][:28] in p.text for p in doc.paragraphs):
            continue
        target = None
        for p in doc.paragraphs:
            if anchor in p.text:
                target = p
                break
        if target is None:
            raise RuntimeError(f"Anchor not found: {anchor}")
        for text in reversed(paragraphs):
            target = insert_after(target, text)
            inserted_count += 1

    if inserted_count:
        doc.save(DOCX)
    print(f"backup={BACKUP.name}")
    print(f"inserted={inserted_count}")


if __name__ == "__main__":
    main()
