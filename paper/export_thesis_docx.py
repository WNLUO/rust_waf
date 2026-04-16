#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
from typing import Iterable

import cairosvg
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


PROJECT_ROOT = Path(__file__).resolve().parent
DEFAULT_HTML = PROJECT_ROOT / "rust_waf_thesis.html"
DEFAULT_OUTPUT = PROJECT_ROOT / "rust_waf_thesis.docx"
DEFAULT_BUILD_DIR = PROJECT_ROOT / ".docx_build"
DEFAULT_TEMPLATE = Path("/Users/wnluo/Downloads/pro/1-基于Spring Boot的反诈宣传平台与实现(1).docx")
DEFAULT_TITLE = "基于Rust的L4/L7协同Web防护系统设计与实现"
DEFAULT_EN_TITLE = "Design and Implementation of an L4/L7 Collaborative Web Protection System Based on Rust"
DEFAULT_DATE = "2026年04月12日"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export the thesis HTML to a template-like DOCX.")
    parser.add_argument("--html", default=str(DEFAULT_HTML), help="Input thesis HTML path")
    parser.add_argument("--out", default=str(DEFAULT_OUTPUT), help="Output DOCX path")
    parser.add_argument(
        "--template",
        default=str(DEFAULT_TEMPLATE) if DEFAULT_TEMPLATE.exists() else "",
        help="Optional DOCX template used as the style/page-layout base",
    )
    parser.add_argument("--build-dir", default=str(DEFAULT_BUILD_DIR), help="Temporary image render directory")
    parser.add_argument("--title", default=DEFAULT_TITLE, help="Chinese thesis title")
    parser.add_argument("--en-title", default=DEFAULT_EN_TITLE, help="English thesis title")
    parser.add_argument("--date", default=DEFAULT_DATE, help="Cover completion date")
    parser.add_argument("--figure-width-cm", type=float, default=15.2, help="Default figure width in cm")
    parser.add_argument(
        "--picture-layout",
        choices=["top-bottom", "inline"],
        default="top-bottom",
        help="DOCX picture layout. top-bottom maps to Word's 上下型环绕.",
    )
    parser.add_argument(
        "--figure-caption-position",
        choices=["below", "above"],
        default="below",
        help="Where to place figure captions in DOCX. Tables are always caption-above.",
    )
    return parser.parse_args()


def set_run_font(
    run,
    *,
    size: float = 12,
    bold: bool = False,
    italic: bool = False,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_common(paragraph, *, indent: bool = True, center: bool = False, line_pt: float = 20) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(line_pt)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if indent and not center:
        paragraph.paragraph_format.first_line_indent = Pt(24)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.gutter = Cm(0)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.50)
    section.different_first_page_header_footer = True
    force_section_layout(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)


def force_section_layout(doc: Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_sz = sect_pr.pgSz
        pg_sz.set(qn("w:w"), "11906")
        pg_sz.set(qn("w:h"), "16838")
        pg_mar = sect_pr.pgMar
        pg_mar.set(qn("w:top"), "1588")
        pg_mar.set(qn("w:right"), "1418")
        pg_mar.set(qn("w:bottom"), "1418")
        pg_mar.set(qn("w:left"), "1701")
        pg_mar.set(qn("w:header"), "850")
        pg_mar.set(qn("w:footer"), "850")
        pg_mar.set(qn("w:gutter"), "0")


def clear_document_keep_section(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=10.5)


def configure_footer(doc: Document) -> None:
    section = doc.sections[0]
    for header in [section.header, section.first_page_header]:
        header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)


def add_blank(doc: Document, count: int = 1, size: float = 12) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        set_paragraph_common(p, indent=False)
        run = p.add_run("")
        set_run_font(run, size=size)


def cover_spacer(doc: Document, line_pt: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(line_pt)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    set_run_font(run, size=1)


def add_cover(doc: Document, *, title: str, date: str) -> None:
    for text in ["校内论文码：______________", "学      号：______________"]:
        p = doc.add_paragraph()
        set_paragraph_common(p, indent=False, center=False, line_pt=20)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, east_asia="黑体")

    cover_spacer(doc, 64)

    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True, line_pt=44)
    run = p.add_run("贵州师范大学本科毕业论文")
    set_run_font(run, size=26, bold=True, east_asia="黑体")

    cover_spacer(doc, 104)

    title_lines = split_cover_title(title)
    title_label_width = "题    目："
    for idx, line in enumerate(title_lines):
        p = doc.add_paragraph()
        set_paragraph_common(p, indent=False, line_pt=24)
        p.paragraph_format.left_indent = Cm(1.65 if idx == 0 else 3.45)
        run = p.add_run(f"{title_label_width}{line}" if idx == 0 else line)
        set_run_font(run, size=15, bold=True, east_asia="黑体")

    cover_spacer(doc, 12)

    meta_lines = [
        "学    院：________________",
        "专    业：________________",
        "年    级：________________",
        "姓    名：________________",
        "指导教师：________________",
        f"完成时间：{date}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        set_paragraph_common(p, indent=False, line_pt=24)
        p.paragraph_format.left_indent = Cm(2.35)
        run = p.add_run(line)
        set_run_font(run, size=15, east_asia="楷体_GB2312")

    doc.add_page_break()


def split_cover_title(title: str, max_units: int = 17) -> list[str]:
    units = 0
    split_at = 0
    preferred_breaks = ["系统", "平台", "设计", "实现", "与"]
    for idx, char in enumerate(title):
        units += 1 if ord(char) > 127 else 0.55
        if units <= max_units:
            split_at = idx + 1
    if split_at >= len(title):
        return [title]

    window = title[:split_at]
    best = 0
    for token in preferred_breaks:
        pos = window.rfind(token)
        if pos > best:
            best = pos + len(token)
    if best >= 8:
        split_at = best

    first = title[:split_at].strip()
    rest = title[split_at:].strip()
    return [first] + split_cover_title(rest, max_units=max_units) if rest else [first]


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True)
    run = p.add_run("目录")
    set_run_font(run, size=16, bold=True, east_asia="黑体")

    p = doc.add_paragraph()
    run_begin = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run_instr._r.append(instr)

    run_sep = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    placeholder = p.add_run("右键目录并选择“更新域”以生成目录")
    set_run_font(placeholder, size=12)

    run_end = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)

    doc.add_page_break()


def add_title_page(doc: Document, *, title: str, en_title: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True)
    run = p.add_run(title)
    set_run_font(run, size=18, east_asia="方正小标宋简体")

    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True)
    run = p.add_run(en_title)
    set_run_font(run, size=14, east_asia="楷体_GB2312")

    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True)
    run = p.add_run("作者：________________")
    set_run_font(run, size=14, east_asia="楷体_GB2312")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 4)}")
    center = level == 1
    set_paragraph_common(paragraph, indent=False, center=center)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, east_asia="黑体")
    elif level == 2:
        set_run_font(run, size=14, bold=True, east_asia="黑体")
    elif level == 3:
        set_run_font(run, size=12, bold=True, east_asia="楷体_GB2312")
    else:
        set_run_font(run, size=12, bold=True)


def add_text_runs(paragraph, node: Tag | NavigableString, *, default_size: float = 12, force_bold: bool = False) -> None:
    if isinstance(node, NavigableString):
        text = str(node).replace("\xa0", " ")
        if text:
            run = paragraph.add_run(text)
            set_run_font(run, size=default_size, bold=force_bold)
        return

    if not isinstance(node, Tag):
        return

    name = node.name.lower()
    if name == "br":
        paragraph.add_run().add_break()
        return

    if name == "sub":
        run = paragraph.add_run(node.get_text("", strip=False))
        set_run_font(run, size=default_size - 1, bold=force_bold)
        run.font.subscript = True
        return

    if name == "sup":
        run = paragraph.add_run(node.get_text("", strip=False))
        set_run_font(run, size=default_size - 1, bold=force_bold)
        run.font.superscript = True
        return

    if name == "code":
        run = paragraph.add_run(node.get_text("", strip=False))
        set_run_font(run, size=10.5, east_asia="宋体", ascii_font="Consolas")
        return

    bold = force_bold or name in {"strong", "b"}
    italic = name in {"em", "i"}
    if italic and node.string:
        run = paragraph.add_run(node.get_text("", strip=False))
        set_run_font(run, size=default_size, bold=bold, italic=True)
        return

    for child in node.children:
        add_text_runs(paragraph, child, default_size=default_size, force_bold=bold)


def add_paragraph_from_tag(doc: Document, tag: Tag, *, center: bool = False, indent: bool = True, bold: bool = False, size: float = 12) -> None:
    p = doc.add_paragraph()
    set_paragraph_common(p, indent=indent, center=center)
    for child in tag.children:
        add_text_runs(p, child, default_size=size, force_bold=bold)
    if not p.runs:
        run = p.add_run(tag.get_text(" ", strip=True))
        set_run_font(run, size=size, bold=bold)


def add_caption(doc: Document, text: str, *, keep_with_next: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True, line_pt=18)
    paragraph_format = p.paragraph_format
    paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True, east_asia="黑体")


def resolve_image(src: str, *, html_path: Path, build_dir: Path) -> Path | None:
    src = src.strip()
    if not src or src.startswith("data:"):
        return None
    image_path = (html_path.parent / src).resolve()
    if not image_path.exists():
        return None
    if image_path.suffix.lower() != ".svg":
        return image_path

    build_dir.mkdir(parents=True, exist_ok=True)
    rendered = build_dir / f"{image_path.stem}.png"
    cairosvg.svg2png(url=str(image_path), write_to=str(rendered), output_width=1800)
    return rendered


def add_picture(doc: Document, path: Path, width_cm: float, *, layout: str = "top-bottom") -> None:
    p = doc.add_paragraph()
    set_paragraph_common(p, indent=False, center=True, line_pt=1 if layout == "top-bottom" else 12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    final_width = fitting_image_width_cm(path, max_width_cm=width_cm, max_height_cm=13.2)
    shape = run.add_picture(str(path), width=Cm(final_width))
    if layout == "top-bottom":
        convert_inline_picture_to_top_bottom(shape._inline)


def convert_inline_picture_to_top_bottom(inline) -> None:
    inline.tag = qn("wp:anchor")
    for key, value in {
        "distT": "36000",
        "distB": "36000",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251659264",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "0",
    }.items():
        inline.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    align_h = OxmlElement("wp:align")
    align_h.text = "center"
    position_h.append(align_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_offset = OxmlElement("wp:posOffset")
    pos_offset.text = "0"
    position_v.append(pos_offset)

    wrap = OxmlElement("wp:wrapTopAndBottom")

    inline.insert(0, simple_pos)
    inline.insert(1, position_h)
    inline.insert(2, position_v)

    # Word expects wrap information before docPr/cNvGraphicFramePr/graphic.
    # The original inline children are usually: extent, docPr, cNvGraphicFramePr, graphic.
    # After inserting simplePos/positionH/positionV, extent is at index 3, so wrap belongs at index 4.
    inline.insert(4, wrap)


def fitting_image_width_cm(path: Path, *, max_width_cm: float, max_height_cm: float) -> float:
    try:
        with Image.open(path) as image:
            width_px, height_px = image.size
    except Exception:
        return max_width_cm
    if width_px <= 0 or height_px <= 0:
        return max_width_cm
    aspect = height_px / width_px
    width_by_height = max_height_cm / aspect
    return min(max_width_cm, width_by_height)


def add_pre(doc: Document, tag: Tag) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(tag.get_text("\n", strip=False))
    set_run_font(run, size=10.5, east_asia="宋体", ascii_font="Courier New")


def add_html_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return
    col_count = max(len(row.find_all(["th", "td"])) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_tag in rows:
        cells = row_tag.find_all(["th", "td"])
        row = table.add_row()
        for idx, cell_tag in enumerate(cells):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(16)
            for child in cell_tag.children:
                add_text_runs(paragraph, child, default_size=10.5, force_bold=(cell_tag.name == "th"))


def add_references(doc: Document, list_tag: Tag) -> None:
    for index, item in enumerate(list_tag.find_all("li", recursive=False), start=1):
        text = item.get_text(" ", strip=True)
        text = re.sub(r"^\[\d+\]\s*", "", text)
        p = doc.add_paragraph()
        set_paragraph_common(p, indent=False)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        run = p.add_run(f"[{index}] {text}")
        set_run_font(run, size=12)


def is_content_node(node: Tag) -> bool:
    if node.name in {"h1", "h2", "h3", "h4", "p", "img", "table", "pre", "ul", "ol"}:
        return True
    classes = node.get("class") or []
    return node.name == "div" and ("title-page" in classes or "page-break" in classes)


def walk(node: Tag) -> Iterable[Tag]:
    if is_content_node(node):
        yield node
        if node.name in {"img", "table", "pre", "ul", "ol"}:
            return
        classes = node.get("class") or []
        if node.name == "div" and ("title-page" in classes or "page-break" in classes):
            return
    for child in node.children:
        if isinstance(child, Tag):
            yield from walk(child)


def iter_main_nodes(soup: BeautifulSoup) -> list[Tag]:
    body = soup.body
    if body is None:
        raise RuntimeError("HTML does not contain <body>.")
    start = body.find("div", class_="title-page")
    started = start is None
    nodes: list[Tag] = []
    for child in body.children:
        if not isinstance(child, Tag):
            continue
        if child is start:
            started = True
        if started:
            nodes.extend(walk(child))
    return nodes


def export_docx(args: argparse.Namespace) -> Path:
    html_path = Path(args.html).resolve()
    output_path = Path(args.out).resolve()
    build_dir = Path(args.build_dir).resolve()
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")

    template_path = Path(args.template).resolve() if args.template else None
    if template_path and template_path.exists():
        doc = Document(str(template_path))
        clear_document_keep_section(doc)
    else:
        doc = Document()
    configure_document(doc)
    configure_footer(doc)
    add_cover(doc, title=args.title, date=args.date)
    add_toc(doc)

    nodes = iter_main_nodes(soup)
    pending_figure_caption: str | None = None
    last_was_page_break = True

    for node in nodes:
        classes = node.get("class") or []
        name = node.name.lower()

        if name == "div" and "title-page" in classes:
            add_title_page(doc, title=args.title, en_title=args.en_title)
            last_was_page_break = False
            continue

        if name == "div" and "page-break" in classes:
            doc.add_page_break()
            last_was_page_break = True
            continue

        if name in {"h1", "h2", "h3", "h4"}:
            if pending_figure_caption:
                add_caption(doc, pending_figure_caption)
                pending_figure_caption = None
            if name == "h1" and not last_was_page_break:
                doc.add_page_break()
            add_heading(doc, node.get_text(" ", strip=True), {"h1": 1, "h2": 2, "h3": 3, "h4": 4}[name])
            last_was_page_break = False
            continue

        if name == "p":
            text = node.get_text(" ", strip=True)
            if not text:
                continue
            if "figure" in classes:
                if args.figure_caption_position == "above":
                    add_caption(doc, text, keep_with_next=True)
                    last_was_page_break = False
                else:
                    pending_figure_caption = text
                continue
            if pending_figure_caption:
                add_caption(doc, pending_figure_caption)
                pending_figure_caption = None
                last_was_page_break = False
            if "table-title" in classes:
                add_caption(doc, text, keep_with_next=True)
                last_was_page_break = False
                continue
            add_paragraph_from_tag(
                doc,
                node,
                center=("center" in classes),
                indent=("no-indent" not in classes and "keywords-line" not in classes),
                size=12,
            )
            last_was_page_break = False
            continue

        if name == "img":
            image_path = resolve_image(node.get("src", ""), html_path=html_path, build_dir=build_dir)
            if image_path:
                add_picture(doc, image_path, args.figure_width_cm, layout=args.picture_layout)
                last_was_page_break = False
            if pending_figure_caption and args.figure_caption_position == "below":
                add_caption(doc, pending_figure_caption)
                pending_figure_caption = None
                last_was_page_break = False
            continue

        if pending_figure_caption:
            add_caption(doc, pending_figure_caption)
            pending_figure_caption = None
            last_was_page_break = False

        if name == "table":
            add_html_table(doc, node)
            last_was_page_break = False
            continue

        if name == "pre":
            add_pre(doc, node)
            last_was_page_break = False
            continue

        if name in {"ul", "ol"} and "refs" in classes:
            add_references(doc, node)
            last_was_page_break = False
            continue

    if pending_figure_caption:
        add_caption(doc, pending_figure_caption)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    cleanup_unused_media(output_path)
    return output_path


def rels_source_path(rels_name: str) -> str | None:
    if not rels_name.endswith(".rels"):
        return None
    if rels_name == "word/_rels/document.xml.rels":
        return "word/document.xml"
    match = re.fullmatch(r"word/_rels/(.+\.xml)\.rels", rels_name)
    if match:
        return f"word/{match.group(1)}"
    return None


def cleanup_unused_media(docx_path: Path) -> None:
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    image_type_suffix = "/image"

    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}

    referenced_media: set[str] = set()

    for rels_name, rels_data in list(entries.items()):
        if not rels_name.startswith("word/_rels/") or not rels_name.endswith(".rels"):
            continue
        source_name = rels_source_path(rels_name)
        source_xml = entries.get(source_name or "")
        if source_xml is None:
            continue

        used_rids = set(re.findall(rb'r:(?:embed|id)="([^"]+)"', source_xml))
        root = ET.fromstring(rels_data)
        changed = False
        for rel in list(root):
            rel_type = rel.attrib.get("Type", "")
            rel_id = rel.attrib.get("Id", "").encode()
            target = rel.attrib.get("Target", "")
            if rel_type.endswith(image_type_suffix) and rel_id not in used_rids:
                root.remove(rel)
                changed = True
                continue
            if rel_type.endswith(image_type_suffix) and target:
                media_name = target.split("/")[-1]
                referenced_media.add(f"word/media/{media_name}")
        if changed:
            ET.register_namespace("", rel_ns)
            entries[rels_name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    entries = {
        name: data
        for name, data in entries.items()
        if not name.startswith("word/media/") or name in referenced_media
    }

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    try:
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for name, data in entries.items():
                zout.writestr(name, data)
        shutil.move(str(tmp_path), docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def main() -> None:
    output_path = export_docx(parse_args())
    print(f"WROTE: {output_path}")


if __name__ == "__main__":
    main()
